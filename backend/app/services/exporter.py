import io
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def export_docx(resume_data: Dict[str, Any]) -> bytes:
    """
    Generates a professional, ATS-friendly DOCX document for the resume.
    """
    doc = Document()
    
    # Page setup - 0.75-inch margins for modern dense text
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styling Helpers
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # 1. Header (Name, Contact Details)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(f"{resume_data['personal_info'].get('first_name', '')} {resume_data['personal_info'].get('last_name', '')}")
    run_name.bold = True
    run_name.font.size = Pt(16)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []
    
    email = resume_data['personal_info'].get('email', '')
    phone = resume_data['personal_info'].get('phone', '')
    location = resume_data['personal_info'].get('location', '')
    github = resume_data['personal_info'].get('github', '')
    linkedin = resume_data['personal_info'].get('linkedin', '')
    website = resume_data['personal_info'].get('website', '')
    
    if email: contact_parts.append(email)
    if phone: contact_parts.append(phone)
    if location: contact_parts.append(location)
    if github: contact_parts.append(f"GitHub: {github}")
    if linkedin: contact_parts.append(f"LinkedIn: {linkedin}")
    if website: contact_parts.append(website)
    
    p_contact.add_run("  |  ".join(contact_parts)).font.size = Pt(9.5)
    
    # Helper to add section headers
    def add_section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11.5)
        # Horizontal divider under section title
        p_border = doc.add_paragraph()
        p_border.paragraph_format.space_before = Pt(0)
        p_border.paragraph_format.space_after = Pt(4)
        run_border = p_border.add_run("―" * 60)
        run_border.font.size = Pt(6)
        run_border.font.color.rgb = RGBColor(128, 128, 128)
        
    # 2. Professional Summary
    summary = resume_data.get('summary', '')
    if summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(6)
        
    # 3. Education
    edu_list = resume_data.get('education', [])
    if edu_list:
        add_section_header("Education")
        for edu in edu_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            inst = p.add_run(edu.get('institution', ''))
            inst.bold = True
            
            p_details = doc.add_paragraph()
            p_details.paragraph_format.space_after = Pt(4)
            deg_field = f"{edu.get('degree', '')} in {edu.get('field_of_study', '')}"
            p_details.add_run(f"{deg_field}  |  GPA: {edu.get('gpa', 'N/A')}")
            
            # Dates
            p_dates = doc.add_paragraph()
            p_dates.paragraph_format.space_after = Pt(6)
            date_range = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
            p_dates.add_run(f"{edu.get('location', '')}  |  {date_range}").italic = True
            
    # 4. Skills & Technologies
    skills = resume_data.get('skills', [])
    techs = resume_data.get('technologies', [])
    if skills or techs:
        add_section_header("Skills & Technologies")
        if skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.add_run("Core Skills: ").bold = True
            p.add_run(", ".join(skills))
        if techs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.add_run("Tools & Technologies: ").bold = True
            p.add_run(", ".join(techs))
            
    # 5. Internships / Experience
    interns = resume_data.get('internships', [])
    if interns:
        add_section_header("Professional Experience")
        for exp in interns:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            comp = p.add_run(exp.get('company', ''))
            comp.bold = True
            p.add_run(f"  ―  {exp.get('role', '')}").italic = True
            
            # Dates/Location
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_after = Pt(4)
            p_meta.add_run(f"{exp.get('location', '')}  |  {exp.get('start_date', '')} - {exp.get('end_date', '')}").italic = True
            
            # Description bullets
            desc = exp.get('description', '')
            for bullet in [b.strip() for b in desc.split('\n') if b.strip()]:
                # Remove starting hyphen if exists
                b_text = bullet[1:].strip() if bullet.startswith('-') or bullet.startswith('•') else bullet
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(b_text)
                
    # 6. Projects
    projs = resume_data.get('projects', [])
    if projs:
        add_section_header("Academic & Personal Projects")
        for proj in projs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            title = p.add_run(proj.get('title', ''))
            title.bold = True
            p.add_run(f"  |  Role: {proj.get('role', 'Developer')}")
            
            if proj.get('technologies'):
                p.add_run(f"  |  ({proj.get('technologies')})").italic = True
                
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_after = Pt(4)
            p_meta.add_run(f"{proj.get('start_date', '')} - {proj.get('end_date', '')}").italic = True
            
            desc = proj.get('description', '')
            for bullet in [b.strip() for b in desc.split('\n') if b.strip()]:
                b_text = bullet[1:].strip() if bullet.startswith('-') or bullet.startswith('•') else bullet
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(b_text)
                
    # 7. Certifications
    certs = resume_data.get('certifications', [])
    if certs:
        add_section_header("Certifications")
        for cert in certs:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            c_run = p.add_run(cert.get('name', ''))
            c_run.bold = True
            p.add_run(f" ― Issued by {cert.get('issuer', 'N/A')} ({cert.get('issue_date', '')})")
            
    # Save to memory stream
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def export_pdf(resume_data: Dict[str, Any]) -> bytes:
    """
    Generates a beautifully typeset, ATS-friendly PDF document.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    style_name = ParagraphStyle(
        'ResumeName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        alignment=1, # Center
        textColor=colors.HexColor('#0f172a') # Slate-900
    )
    
    style_contact = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        alignment=1, # Center
        textColor=colors.HexColor('#475569') # Slate-600
    )
    
    style_section_title = ParagraphStyle(
        'ResumeSectionTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#1e3a8a'), # Indigo-900
        spaceBefore=10,
        keepWithNext=True
    )
    
    style_body = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#1e293b') # Slate-800
    )
    
    style_bullet = ParagraphStyle(
        'ResumeBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        leftIndent=15,
        firstLineIndent=-10,
        textColor=colors.HexColor('#334155'), # Slate-700
        spaceBefore=2
    )
    
    story = []
    
    # 1. Header (Name & Contact)
    personal = resume_data['personal_info']
    story.append(Paragraph(f"{personal.get('first_name', '')} {personal.get('last_name', '')}", style_name))
    story.append(Spacer(1, 4))
    
    contact_parts = []
    email = personal.get('email', '')
    phone = personal.get('phone', '')
    location = personal.get('location', '')
    github = personal.get('github', '')
    linkedin = personal.get('linkedin', '')
    website = personal.get('website', '')
    
    if email: contact_parts.append(email)
    if phone: contact_parts.append(phone)
    if location: contact_parts.append(location)
    if github: contact_parts.append(f"GitHub: {github}")
    if linkedin: contact_parts.append(f"LinkedIn: {linkedin}")
    if website: contact_parts.append(website)
    
    story.append(Paragraph(" &nbsp;|&nbsp; ".join(contact_parts), style_contact))
    story.append(Spacer(1, 10))
    
    # Section Header Helper
    def build_section(title: str):
        story.append(Paragraph(title.upper(), style_section_title))
        # Draw a clean horizontal line
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'), spaceBefore=2, spaceAfter=6))
        
    # 2. Professional Summary
    summary = resume_data.get('summary', '')
    if summary:
        build_section("Professional Summary")
        story.append(Paragraph(summary, style_body))
        
    # 3. Education
    edu_list = resume_data.get('education', [])
    if edu_list:
        build_section("Education")
        for edu in edu_list:
            degree_str = f"<b>{edu.get('degree', '')} in {edu.get('field_of_study', '')}</b>"
            inst_str = f"<i>{edu.get('institution', '')}</i>"
            gpa_str = f"GPA: {edu.get('gpa', 'N/A')}"
            dates_str = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
            loc_str = edu.get('location', '')
            
            # Use Table for alignment
            data = [
                [Paragraph(degree_str, style_body), Paragraph(dates_str, ParagraphStyle('RightText', parent=style_body, alignment=2))],
                [Paragraph(f"{inst_str} ({gpa_str})", style_body), Paragraph(loc_str, ParagraphStyle('RightText', parent=style_body, alignment=2))]
            ]
            t = Table(data, colWidths=[360, 180])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(t)
            story.append(Spacer(1, 4))
            
    # 4. Skills & Technologies
    skills = resume_data.get('skills', [])
    techs = resume_data.get('technologies', [])
    if skills or techs:
        build_section("Skills & Technologies")
        if skills:
            story.append(Paragraph(f"<b>Core Skills:</b> {', '.join(skills)}", style_body))
            story.append(Spacer(1, 3))
        if techs:
            story.append(Paragraph(f"<b>Tools & Technologies:</b> {', '.join(techs)}", style_body))
            story.append(Spacer(1, 4))
            
    # 5. Internships / Experience
    interns = resume_data.get('internships', [])
    if interns:
        build_section("Professional Experience")
        for exp in interns:
            comp_role = f"<b>{exp.get('company', '')}</b>  ―  <i>{exp.get('role', '')}</i>"
            dates_str = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
            loc_str = exp.get('location', '')
            
            data = [
                [Paragraph(comp_role, style_body), Paragraph(dates_str, ParagraphStyle('RightText', parent=style_body, alignment=2))],
                [Paragraph("", style_body), Paragraph(loc_str, ParagraphStyle('RightText', parent=style_body, alignment=2))]
            ]
            t = Table(data, colWidths=[360, 180])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(t)
            
            # Bullets
            desc = exp.get('description', '')
            for bullet in [b.strip() for b in desc.split('\n') if b.strip()]:
                b_text = bullet[1:].strip() if bullet.startswith('-') or bullet.startswith('•') else bullet
                story.append(Paragraph(f"&bull;&nbsp;&nbsp;{b_text}", style_bullet))
            story.append(Spacer(1, 4))
            
    # 6. Projects
    projs = resume_data.get('projects', [])
    if projs:
        build_section("Academic & Personal Projects")
        for proj in projs:
            title_role = f"<b>{proj.get('title', '')}</b>  |  Role: {proj.get('role', 'Developer')}"
            dates_str = f"{proj.get('start_date', '')} - {proj.get('end_date', '')}"
            tech_str = f"({proj.get('technologies', '')})" if proj.get('technologies') else ""
            
            data = [
                [Paragraph(title_role, style_body), Paragraph(dates_str, ParagraphStyle('RightText', parent=style_body, alignment=2))],
                [Paragraph(f"<i>{tech_str}</i>", style_body), Paragraph("", style_body)]
            ]
            t = Table(data, colWidths=[360, 180])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(t)
            
            desc = proj.get('description', '')
            for bullet in [b.strip() for b in desc.split('\n') if b.strip()]:
                b_text = bullet[1:].strip() if bullet.startswith('-') or bullet.startswith('•') else bullet
                story.append(Paragraph(f"&bull;&nbsp;&nbsp;{b_text}", style_bullet))
            story.append(Spacer(1, 4))
            
    # 7. Certifications
    certs = resume_data.get('certifications', [])
    if certs:
        build_section("Certifications")
        for cert in certs:
            cert_text = f"<b>{cert.get('name', '')}</b> ― Issued by {cert.get('issuer', 'N/A')} ({cert.get('issue_date', '')})"
            story.append(Paragraph(f"&bull;&nbsp;&nbsp;{cert_text}", style_bullet))
            
    # Build document
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes
